from __future__ import annotations

from datetime import date
from pathlib import Path
import hashlib

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "Research Assets" / "Documentation"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "Gates Manfluencer Project Research Handover and Dataset Guide.docx"


def file_md5(path: Path) -> str:
    h = hashlib.md5()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def excel_row_count(path: Path, all_sheets: bool = False) -> int:
    try:
        if not all_sheets:
            return int(len(pd.read_excel(path)))
        xls = pd.ExcelFile(path)
        return int(sum(len(pd.read_excel(path, sheet_name=sn)) for sn in xls.sheet_names))
    except Exception:
        return -1


def normalize(text: str) -> str:
    return str(text).lower().replace("\n", " ").strip()


def main() -> int:
    sample_dir = ROOT / "Research Assets" / "Sample Influencer Lists"
    kenya_list = pd.read_excel(sample_dir / "Kenya Manfluencer List.xlsx", sheet_name="Top 6")
    nigeria_list = pd.read_excel(sample_dir / "Nigeria Manfluencer Lists.xlsx", sheet_name="Top 6")
    nigeria_examples = pd.read_excel(sample_dir / "Nigeria Content Examples.xlsx", sheet_name="Nigeria")

    creators = [
        {"country": "Kenya", "name": "Eric Amunga (Amerix)", "orientation": "Regressive", "primary_platform": "X"},
        {"country": "Kenya", "name": "Andrew Kibe", "orientation": "Regressive", "primary_platform": "YouTube / X"},
        {"country": "Kenya", "name": "Philip Karanja", "orientation": "Progressive", "primary_platform": "YouTube"},
        {"country": "Kenya", "name": "Onyango Otieno (Rixpoet)", "orientation": "Progressive", "primary_platform": "YouTube"},
        {"country": "Kenya", "name": "Eddy Kimani", "orientation": "Progressive", "primary_platform": "Instagram / X"},
        {"country": "Nigeria", "name": "Banky Wellington", "orientation": "Progressive", "primary_platform": "YouTube / Instagram / X"},
        {"country": "Nigeria", "name": "Deyemi Okanlawon", "orientation": "Progressive", "primary_platform": "X"},
        {"country": "Nigeria", "name": "Wizarab", "orientation": "Regressive", "primary_platform": "X"},
        {"country": "Nigeria", "name": "Shola", "orientation": "Regressive", "primary_platform": "X"},
        {"country": "Nigeria", "name": "Agba John Doe", "orientation": "Regressive", "primary_platform": "X"},
    ]

    def find_reach_row(name: str):
        needle = normalize(name).split("(")[0].strip()[:12]
        for df, src in [
            (kenya_list, "Kenya Manfluencer List.xlsx::Top 6"),
            (nigeria_list, "Nigeria Manfluencer Lists.xlsx::Top 6"),
        ]:
            if "Influencer Name" not in df.columns:
                continue
            hits = df[df["Influencer Name"].fillna("").astype(str).str.lower().str.contains(needle, regex=False)]
            if len(hits):
                row = hits.iloc[0]
                reach = "; ".join(
                    [
                        f"YT: {row.get('YouTube Subs (as-of)', '')}",
                        f"IG: {row.get('IG Followers (as-of)', '')}",
                        f"TikTok: {row.get('TikTok Followers (as-of)', '')}",
                        f"X: {row.get('X Followers (as-of)', '')}",
                    ]
                )
                return reach, src, str(row.get("Source Links", "")).strip()

        if "Influencer" in nigeria_examples.columns:
            hits = nigeria_examples[
                nigeria_examples["Influencer"].fillna("").astype(str).str.lower().str.contains(needle[:10], regex=False)
            ]
            if len(hits):
                row = hits.iloc[0]
                reach = str(row.get("Social Handles & Reach", "")).replace("\n", "; ")
                why = str(row.get("Why Included", "")).replace("\n", " ")
                return reach, "Nigeria Content Examples.xlsx::Nigeria", why
        return "Not found in candidate sheets", "N/A", "N/A"

    for creator in creators:
        reach, source, why = find_reach_row(creator["name"])
        creator["reach_evidence"] = reach
        creator["evidence_source"] = source
        creator["selection_rationale_note"] = why

    inventory: list[dict] = []

    for p in sorted((ROOT / "Kenya" / "Audience Comments - Raw").glob("*.xlsx")):
        inventory.append({"tier": "Kenya Audience Raw", "path": p, "rows": excel_row_count(p), "hash": file_md5(p)})
    for p in sorted((ROOT / "Kenya" / "Audience Comments - Filtered").glob("*.xlsx")):
        inventory.append({"tier": "Kenya Audience Filtered", "path": p, "rows": excel_row_count(p), "hash": file_md5(p)})
    for p in [
        ROOT / "Kenya" / "Audience Analysis" / "Kenya Audience Analysis Final.xlsx",
        ROOT / "Kenya" / "Content Analysis" / "Kenya Content Analysis Final.xlsx",
    ]:
        if p.exists():
            inventory.append({"tier": "Kenya Final Workbook", "path": p, "rows": excel_row_count(p, all_sheets=True), "hash": file_md5(p)})

    for tier, subdir in [
        ("Nigeria Audience Raw", ROOT / "Nigeria" / "Audience Analysis" / "Audience Comments - Raw"),
        ("Nigeria Audience Complete", ROOT / "Nigeria" / "Audience Analysis" / "Audience Comments - Complete"),
        ("Nigeria Audience Final", ROOT / "Nigeria" / "Audience Analysis" / "Audience Comments - Final"),
    ]:
        for p in sorted(subdir.rglob("*.xlsx")):
            inventory.append({"tier": tier, "path": p, "rows": excel_row_count(p), "hash": file_md5(p)})

    for tier, subdir in [
        ("Nigeria Content Raw", ROOT / "Nigeria" / "Content Analysis" / "Content - Raw"),
        ("Nigeria Content Final", ROOT / "Nigeria" / "Content Analysis" / "Content - Final"),
    ]:
        for p in sorted(subdir.rglob("*.xlsx")):
            inventory.append({"tier": tier, "path": p, "rows": excel_row_count(p), "hash": file_md5(p)})

    for p in [
        ROOT / "Nigeria" / "Audience Analysis" / "Nigeria Audience Analysis Final.xlsx",
        ROOT / "Nigeria" / "Content Analysis" / "Nigeria Content Analysis Final.xlsx",
    ]:
        if p.exists():
            inventory.append({"tier": "Nigeria Final Workbook", "path": p, "rows": excel_row_count(p, all_sheets=True), "hash": file_md5(p)})

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run("GATES Manfluencer Project\nResearch Handover and Dataset Governance Guide")
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph(
        f"Version 1.0 | Generated: {date.today().isoformat()} | Repository Root: {ROOT}"
    )
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading("1. Purpose and Scope", level=1)
    doc.add_paragraph(
        "This document is the operational handover for future researchers. It records dataset lineage, canonical "
        "file locations, influencer selection logic, quality controls, and maintenance procedures."
    )

    doc.add_heading("2. Research Questions and Analytical Intent", level=1)
    doc.add_paragraph(
        "The project analyzes masculinity narratives and audience reception across Kenya and Nigeria, with explicit "
        "contrast between progressive and regressive framing."
    )

    doc.add_heading("3. Influencer Selection Framework (Reach + Relevance)", level=1)
    for criterion in [
        "Reach and visibility (follower/subscriber scale and engagement signal).",
        "Masculinity topical fit (gender roles, relationships, fatherhood, emotional norms, anti-toxic narratives).",
        "Ideological coverage (progressive and regressive cohorts).",
        "Platform fit for observable audience discourse (especially X and YouTube comments/replies).",
        "Corpus usability (sufficient post/comment density and traceable source links).",
    ]:
        doc.add_paragraph(criterion, style="List Bullet")

    doc.add_heading("4. Selected Influencers: Canonical Cohort and Evidence", level=1)
    table = doc.add_table(rows=1, cols=8)
    for i, head in enumerate(
        [
            "Country",
            "Influencer",
            "Orientation",
            "Primary Platform",
            "Reach Evidence",
            "Evidence Source",
            "Selection Rationale",
            "Maintenance Note",
        ]
    ):
        table.rows[0].cells[i].text = head

    for creator in creators:
        row = table.add_row().cells
        row[0].text = creator["country"]
        row[1].text = creator["name"]
        row[2].text = creator["orientation"]
        row[3].text = creator["primary_platform"]
        row[4].text = str(creator["reach_evidence"])
        row[5].text = str(creator["evidence_source"])
        if creator["selection_rationale_note"] in {"N/A", "nan"}:
            row[6].text = (
                "Included to preserve country and orientation balance while retaining high-signal masculinity discourse."
            )
        else:
            row[6].text = str(creator["selection_rationale_note"])
        row[7].text = "If roster changes, document reach delta and maintain orientation/country balance."

    doc.add_heading("5. Canonical Dataset Architecture", level=1)
    for path in [
        "Kenya/Audience Comments - Raw/",
        "Kenya/Audience Comments - Filtered/",
        "Kenya/Audience Analysis/Kenya Audience Analysis Final.xlsx",
        "Kenya/Content Analysis/Kenya Content Analysis Final.xlsx",
        "Nigeria/Audience Analysis/Audience Comments - Raw/",
        "Nigeria/Audience Analysis/Audience Comments - Complete/",
        "Nigeria/Audience Analysis/Audience Comments - Final/",
        "Nigeria/Audience Analysis/Nigeria Audience Analysis Final.xlsx",
        "Nigeria/Content Analysis/Content - Raw/",
        "Nigeria/Content Analysis/Content - Final/",
        "Nigeria/Content Analysis/Nigeria Content Analysis Final.xlsx",
        "Research Assets/Sample Influencer Lists/",
    ]:
        doc.add_paragraph(path, style="List Bullet")

    doc.add_heading("6. Data Dictionary (Core Sheets)", level=1)
    doc.add_paragraph(
        "Audience sheets typically include: Comment ID, Influencer, Platform, Source URL, Comment. "
        "Content sheets typically include: Content/Segment ID, Influencer, Platform, Content Type, Source URL, Context, Verbatim Text."
    )
    doc.add_paragraph(
        "Tier semantics: Raw = minimally transformed source capture; Complete = cleaned/normalized broad set; "
        "Filtered/Final = scope-relevant or locked analytical subset used for reporting."
    )

    doc.add_heading("7. Inventory and Provenance Register", level=1)
    doc.add_paragraph(
        "Each file below includes tier, row count, and MD5 fingerprint for reproducibility and post-handoff verification."
    )
    inv = doc.add_table(rows=1, cols=5)
    for i, head in enumerate(["Tier", "Relative Path", "Rows", "MD5", "Status"]):
        inv.rows[0].cells[i].text = head
    for item in sorted(inventory, key=lambda x: (x["tier"], str(x["path"]))):
        row = inv.add_row().cells
        row[0].text = item["tier"]
        row[1].text = str(item["path"].relative_to(ROOT))
        row[2].text = str(item["rows"])
        row[3].text = item["hash"]
        row[4].text = "Canonical"

    doc.add_heading("8. Quality Assurance and De-duplication Protocol", level=1)
    for point in [
        "No duplicate files by content hash are permitted in active dataset areas.",
        "Keep one canonical copy per dataset in operational country folders.",
        "Any update to final workbooks must update this guide (row counts + hashes).",
        "Do not overwrite Raw with transformed data; transformations must remain tiered.",
        "Preserve Source URL fields to keep traceability/auditability.",
    ]:
        doc.add_paragraph(point, style="List Bullet")

    doc.add_heading("9. Maintenance SOP for Future Researchers", level=1)
    for step in [
        "Ingest new data into a staging area and run dedup (MD5) against canonical paths.",
        "Promote only validated files into canonical country folders.",
        "Recompute row counts + hashes and append changelog.",
        "Update influencer evidence if roster changes.",
        "Maintain orientation and country balance unless protocol change is approved.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("10. Changelog (Use for Every Dataset Revision)", level=1)
    chg = doc.add_table(rows=2, cols=5)
    for i, head in enumerate(["Date", "Editor", "Files Added/Removed/Updated", "Reason", "Integrity Notes (row/hash)"]):
        chg.rows[0].cells[i].text = head
    chg.rows[1].cells[0].text = date.today().isoformat()
    chg.rows[1].cells[1].text = "Initial handover generation"
    chg.rows[1].cells[2].text = "Baseline register generated"
    chg.rows[1].cells[3].text = "Project continuity and reproducibility"
    chg.rows[1].cells[4].text = "Canonical inventory captured"

    doc.add_heading("11. Key Reference Files", level=1)
    for ref in [
        "README.md",
        "Research Assets/Sample Influencer Lists/Gates_ Positive Masculinity Influencers by Social Media Metrics.xlsx",
        "Research Assets/Sample Influencer Lists/Kenya Manfluencer List.xlsx",
        "Research Assets/Sample Influencer Lists/Nigeria Manfluencer Lists.xlsx",
        "Research Assets/Sample Influencer Lists/Nigeria Content Examples.xlsx",
    ]:
        doc.add_paragraph(ref, style="List Bullet")

    doc.save(OUT_PATH)
    print(f"WROTE {OUT_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
