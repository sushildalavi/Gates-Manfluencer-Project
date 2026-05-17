from __future__ import annotations

import re
from pathlib import Path
from collections import Counter, defaultdict

import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


BASE = Path('/Users/sushildalavi/Desktop/NLC/Gates-Manfluencer-Project')
OUT_DOC = BASE / 'Research Assets/Deliverables/LLM Exploratory Analysis Revised.docx'
FIG_DIR = BASE / 'Research Assets/Deliverables/Figures'
FIG_DIR.mkdir(parents=True, exist_ok=True)


def clean_text(x):
    if pd.isna(x):
        return ''
    s = str(x).strip()
    return '' if s.lower() in {'nan', 'none'} else s


def split_multi(value: str) -> list[str]:
    s = clean_text(value)
    if not s:
        return []
    s = re.sub(r'\s+/\s+', ' | ', s)
    s = s.replace('\n', ' | ')
    parts = [p.strip(' -•\t') for p in re.split(r'\||,|;|/|\\n', s) if p.strip()]
    return [p for p in parts if p and p.lower() not in {'not applicable', 'n/a', 'na'}]


def pct(n, d):
    if not d:
        return '0.0%'
    return f'{(100.0 * n / d):.1f}%'


def read_final_counts(path: Path, summary_names: set[str]):
    xl = pd.ExcelFile(path)
    out = {}
    for sheet in xl.sheet_names:
        if sheet in summary_names:
            continue
        out[sheet] = len(pd.read_excel(path, sheet_name=sheet))
    return out


def short_name_from_sheet(sheet: str):
    s = sheet
    for marker in ['(', '_', ' - ', ' I wonder', ' My voice', ' Men are', ' A woman']:
        if marker in s:
            s = s.split(marker)[0]
            break
    return s.strip()


def make_charts(final_counts, orientation_counts, metric_presence):
    # Chart 1: final sample sizes
    labels = ['Kenya Content', 'Kenya Audience', 'Nigeria Content', 'Nigeria Audience']
    values = [
        sum(final_counts['kenya_content'].values()),
        sum(final_counts['kenya_audience'].values()),
        sum(final_counts['nigeria_content'].values()),
        sum(final_counts['nigeria_audience'].values()),
    ]
    plt.figure(figsize=(8, 4.5))
    bars = plt.bar(labels, values, color=['#2F6B3D', '#5A9E6F', '#8A4F7D', '#B370A0'])
    plt.title('Final Rows Included in Current Exploratory Pack')
    plt.ylabel('Rows')
    plt.xticks(rotation=10)
    for b, v in zip(bars, values):
        plt.text(b.get_x() + b.get_width() / 2, v, str(v), ha='center', va='bottom', fontsize=9)
    plt.tight_layout()
    p1 = FIG_DIR / 'sample composition by stream.png'
    plt.savefig(p1, dpi=180)
    plt.close()

    # Chart 2: orientation balance
    x_labels = ['Kenya Content', 'Kenya Audience', 'Nigeria Content', 'Nigeria Audience']
    reg = [
        orientation_counts['kenya_content'].get('Regressive', 0),
        orientation_counts['kenya_audience'].get('Regressive', 0),
        orientation_counts['nigeria_content'].get('Regressive', 0),
        orientation_counts['nigeria_audience'].get('Regressive', 0),
    ]
    pro = [
        orientation_counts['kenya_content'].get('Progressive', 0),
        orientation_counts['kenya_audience'].get('Progressive', 0),
        orientation_counts['nigeria_content'].get('Progressive', 0),
        orientation_counts['nigeria_audience'].get('Progressive', 0),
    ]
    idx = range(len(x_labels))
    w = 0.38
    plt.figure(figsize=(8, 4.5))
    plt.bar([i - w/2 for i in idx], reg, width=w, label='Regressive coded creators', color='#C44E52')
    plt.bar([i + w/2 for i in idx], pro, width=w, label='Progressive coded creators', color='#4C72B0')
    plt.title('Rows by Orientation Group in Final Workbooks')
    plt.ylabel('Rows')
    plt.xticks(list(idx), x_labels, rotation=10)
    plt.legend()
    plt.tight_layout()
    p2 = FIG_DIR / 'orientation balance by stream.png'
    plt.savefig(p2, dpi=180)
    plt.close()

    # Chart 3: metric availability by country
    fields = ['likes', 'reply_count', 'retweets', 'views', 'impressions', 'shares']
    kc = [metric_presence['Kenya'].get(f, 0) for f in fields]
    nc = [metric_presence['Nigeria'].get(f, 0) for f in fields]
    plt.figure(figsize=(8, 4.5))
    ix = range(len(fields))
    plt.bar([i - w/2 for i in ix], kc, width=w, label='Kenya raw audience files', color='#3C8D40')
    plt.bar([i + w/2 for i in ix], nc, width=w, label='Nigeria raw audience files', color='#7D4E9F')
    plt.title('Metric Fields Present in Raw Audience Files')
    plt.ylabel('Number of files with field present')
    plt.xticks(list(ix), fields)
    plt.legend()
    plt.tight_layout()
    p3 = FIG_DIR / 'metric availability by country.png'
    plt.savefig(p3, dpi=180)
    plt.close()

    return [p1, p2, p3]


def metric_field_presence(folder: Path):
    metric_map = {
        'likes': ['likes', 'like_count', 'favorite_count'],
        'reply_count': ['reply_count', 'replies', 'reply'],
        'retweets': ['retweets', 'retweet_count', 'reposts', 'repost_count'],
        'views': ['views', 'view_count'],
        'impressions': ['impressions', 'impression_count'],
        'shares': ['share', 'share_count'],
    }
    files = sorted(folder.rglob('*.xlsx'))
    files = [p for p in files if '/Archive/' not in str(p)]
    field_file_count = Counter()

    for p in files:
        try:
            df = pd.read_excel(p)
        except Exception:
            continue
        cols = [str(c).strip().lower() for c in df.columns]
        for out_field, aliases in metric_map.items():
            if any(any(alias in c for alias in aliases) for c in cols):
                field_file_count[out_field] += 1

    return len(files), dict(field_file_count)


def load_reach_rows(path: Path, sheet: str, selected_names: list[str], total_col: str):
    df = pd.read_excel(path, sheet_name=sheet)
    rows = []
    for name in selected_names:
        candidates = df[df['Influencer Name'].astype(str).str.contains(name, case=False, na=False)]
        if candidates.empty:
            rows.append({'name': name, 'found': False})
            continue
        r = candidates.iloc[0]
        rows.append(
            {
                'name': clean_text(r.get('Influencer Name', name)),
                'found': True,
                'orientation': clean_text(r.get('Orientation', '')),
                'youtube': clean_text(r.get('YouTube Subs (as-of)', '')),
                'instagram': clean_text(r.get('IG Followers (as-of)', '')),
                'tiktok': clean_text(r.get('TikTok Followers (as-of)', '')),
                'x': clean_text(r.get('X Followers (as-of)', '')),
                'total': clean_text(r.get(total_col, '')),
                'reach_tier': clean_text(r.get('Reach Tier', '')),
            }
        )
    return rows


def parse_llm_country(df: pd.DataFrame, country_name: str):
    # These column lookups are resilient to long prompt text prefixes.
    col_topic = next(c for c in df.columns if str(c).startswith('Q2. What is/are the primary topic'))
    col_orientation = next(c for c in df.columns if str(c).startswith('Q5. If yes, how would you characterize the type of masculinity'))
    col_content_type = next(c for c in df.columns if str(c).startswith('Q3. Characterize the type of content'))

    topic_ctr = Counter()
    for v in df[col_topic].fillna(''):
        for p in split_multi(v):
            topic_ctr[p] += 1

    orientation_ctr = Counter(clean_text(v) for v in df[col_orientation].fillna('') if clean_text(v))
    type_ctr = Counter(clean_text(v) for v in df[col_content_type].fillna('') if clean_text(v))

    total = len(df)
    return {
        'country': country_name,
        'n': total,
        'top_topics': topic_ctr.most_common(5),
        'orientation_dist': orientation_ctr,
        'content_type_dist': type_ctr,
    }


def parse_audience_country(df: pd.DataFrame):
    col_sent = next(c for c in df.columns if str(c).startswith('Q1. Overall sentiment'))
    col_stance = next(c for c in df.columns if str(c).startswith('Q8. What is the commenter’s stance'))
    col_topic = next(c for c in df.columns if str(c).startswith('Q7. What is the main topic of the comment'))

    sent = Counter(clean_text(v) for v in df[col_sent].fillna('') if clean_text(v))
    stance = Counter(clean_text(v) for v in df[col_stance].fillna('') if clean_text(v))
    topic = Counter()
    for v in df[col_topic].fillna(''):
        for p in split_multi(v):
            topic[p] += 1
    return {
        'n': len(df),
        'sentiment': sent,
        'stance': stance,
        'topics': topic.most_common(5),
    }


def add_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r in rows:
        cells = table.add_row().cells
        for i, c in enumerate(r):
            cells[i].text = str(c)
    return table


def main():
    # Final workbook row counts
    kenya_audience = read_final_counts(BASE / 'Kenya/Audience Analysis/Kenya Audience Analysis Final.xlsx', {'Summary Metrics'})
    kenya_content = read_final_counts(BASE / 'Kenya/Content Analysis/Kenya Content Analysis Final.xlsx', {'Summary'})
    nigeria_audience = read_final_counts(BASE / 'Nigeria/Audience Analysis/Nigeria Audience Analysis Final.xlsx', {'Summary and Stats'})
    nigeria_content = read_final_counts(BASE / 'Nigeria/Content Analysis/Nigeria Content Analysis Final.xlsx', {'Summary and Stats'})

    # Orientation mapping used in current project package
    kenya_content_orientation_map = {
        'Andrew Kibe': 'Regressive',
        'Eric (Amerix)': 'Regressive',
        'Eddy Kimani': 'Progressive',
        'Onyango Otieno (Rixpoet)': 'Progressive',
        'Philip Karanja': 'Progressive',
    }
    kenya_audience_orientation_map = {
        'Andrew I wonder how men': 'Regressive',
        'Eric A woman cant love': 'Regressive',
        'Eddy Men are evolving': 'Progressive',
        'Rixpoet My voice was beaten': 'Progressive',
    }
    nigeria_content_orientation_map = {
        'Agba John Doe': 'Regressive',
        'Shola': 'Regressive',
        'Wizarab': 'Regressive',
        'Banky Wellington': 'Progressive',
        'Deyemi Okanlawon': 'Progressive',
        'Ebuka Obi-Uchendu': 'Progressive',
    }
    nigeria_audience_orientation_map = {
        'Agba John Doe': 'Regressive',
        'Shola': 'Regressive',
        'Banky Wellington': 'Progressive',
        'Deyemi Okanlawon': 'Progressive',
    }

    def group_orientation(counts, mapping):
        d = defaultdict(int)
        for k, v in counts.items():
            if k in mapping:
                d[mapping[k]] += int(v)
        return dict(d)

    final_counts = {
        'kenya_content': kenya_content,
        'kenya_audience': kenya_audience,
        'nigeria_content': nigeria_content,
        'nigeria_audience': nigeria_audience,
    }
    orientation_counts = {
        'kenya_content': group_orientation(kenya_content, kenya_content_orientation_map),
        'kenya_audience': group_orientation(kenya_audience, kenya_audience_orientation_map),
        'nigeria_content': group_orientation(nigeria_content, nigeria_content_orientation_map),
        'nigeria_audience': group_orientation(nigeria_audience, nigeria_audience_orientation_map),
    }

    # Metric field availability
    kenya_raw_n, kenya_metric_presence = metric_field_presence(BASE / 'Kenya/Audience Comments - Raw')
    nigeria_raw_n, nigeria_metric_presence = metric_field_presence(BASE / 'Nigeria/Audience Analysis/Audience Comments - Raw')
    metric_presence = {'Kenya': kenya_metric_presence, 'Nigeria': nigeria_metric_presence}

    # Reach rows from longlists
    kenya_reach = load_reach_rows(
        BASE / 'Research Assets/Sample Influencer Lists/Kenya Manfluencer List.xlsx',
        'Kenya Longlist',
        ['Andrew Kibe', 'Amerix', 'Eddy Kimani', 'Rixpoet', 'Philip Karanja'],
        'Total Cross-Platform Reach',
    )
    nigeria_reach = load_reach_rows(
        BASE / 'Research Assets/Sample Influencer Lists/Nigeria Manfluencer Lists.xlsx',
        'Nigeria Longlist',
        ['Agba John Doe', 'Banky Wellington', 'Deyemi Okanlawon', 'Shola', 'Wizarab', 'Ebuka Obi-Uchendu'],
        'Total Known Platform Reach',
    )

    # LLM exploratory codebooks (sample-based)
    content_book = BASE / 'Codebooks/LLM Codebook/LLM Coding - Content Analysis.xlsx'
    audience_book = BASE / 'Codebooks/LLM Codebook/LLM Coding - Audience Analysis.xlsx'
    k_content_llm = pd.read_excel(content_book, sheet_name='Kenya - LLM Coding')
    n_content_llm = pd.read_excel(content_book, sheet_name='Nigeria - LLM Coding')
    k_audience_llm = pd.read_excel(audience_book, sheet_name='Kenya - LLM Coding')
    n_audience_llm = pd.read_excel(audience_book, sheet_name='Nigeria - LLM Coding')

    k_content_llm_stats = parse_llm_country(k_content_llm, 'Kenya')
    n_content_llm_stats = parse_llm_country(n_content_llm, 'Nigeria')
    k_audience_llm_stats = parse_audience_country(k_audience_llm)
    n_audience_llm_stats = parse_audience_country(n_audience_llm)

    charts = make_charts(final_counts, orientation_counts, metric_presence)

    # Build document
    doc = Document()
    doc.add_heading('LLM Exploratory Analysis Revised', 0)
    doc.add_paragraph('Date: 2026-05-17')
    doc.add_paragraph('Project: Gates Manfluencer Research (Kenya and Nigeria)')

    doc.add_heading('Purpose of this revision', level=1)
    doc.add_paragraph(
        'This revision addresses review comments requesting clearer methods context, explicit definitions, normalized interpretation, country-specific reporting, and an engagement/reach layer. '
        'All numbers below are drawn from current files in this repository and are labeled as exploratory where based on sampled LLM coding outputs.'
    )

    doc.add_heading('Scope and methods context', level=1)
    doc.add_paragraph('Exploratory versus focused analysis')
    doc.add_paragraph(
        'This document is exploratory only. It summarizes directional findings and data readiness from current coded samples and final workbooks. '
        'It is not a causal or population-level estimate and should not be treated as final focused analysis.'
    )
    doc.add_paragraph('How creators were selected and why this matters for interpretation')
    doc.add_paragraph(
        'Creators were purposively selected to include both progressive-leaning and regressive-leaning masculinity voices, not randomly sampled from all creators in each country. '
        'As a result, percentages reflect this selected corpus and cannot be interpreted as country prevalence estimates.'
    )
    doc.add_paragraph('Definitions used in this exploratory pack')
    doc.add_paragraph(
        'Regressive coding indicates content that reinforces hierarchy, domination, rigid gender roles, female submission, anti-feminist frames, or male grievance as a dominant lens. '
        'Progressive coding indicates content that supports equality, partnership, emotional openness, shared care, non-violence, and expanded masculinities. '
        'Mixed or unclear coding is used where signals conflict or are not explicit.'
    )
    doc.add_paragraph('Comparability note')
    doc.add_paragraph(
        'Kenya and Nigeria are reported separately in this revision. Cross-country direct comparison is methodologically limited because of different creator mixes, platform distributions, and available engagement metadata.'
    )

    doc.add_heading('Current data inventory and final row counts', level=1)
    add_table(
        doc,
        ['Country', 'Stream', 'Final rows', 'Creators in final workbook'],
        [
            ['Kenya', 'Content', sum(kenya_content.values()), len(kenya_content)],
            ['Kenya', 'Audience', sum(kenya_audience.values()), len(kenya_audience)],
            ['Nigeria', 'Content', sum(nigeria_content.values()), len(nigeria_content)],
            ['Nigeria', 'Audience', sum(nigeria_audience.values()), len(nigeria_audience)],
        ],
    )

    doc.add_paragraph('Figure: Current final row volumes used in this revision')
    doc.add_picture(str(charts[0]), width=Inches(6.7))

    doc.add_heading('Normalization and orientation balance', level=1)
    add_table(
        doc,
        ['Stream', 'Regressive rows', 'Progressive rows', 'Interpretation note'],
        [
            [
                'Kenya Content',
                orientation_counts['kenya_content'].get('Regressive', 0),
                orientation_counts['kenya_content'].get('Progressive', 0),
                'Five selected creators; row totals should be read within this curated sample.',
            ],
            [
                'Kenya Audience',
                orientation_counts['kenya_audience'].get('Regressive', 0),
                orientation_counts['kenya_audience'].get('Progressive', 0),
                'Audience rows mirror selected source posts and filtering pipeline.',
            ],
            [
                'Nigeria Content',
                orientation_counts['nigeria_content'].get('Regressive', 0),
                orientation_counts['nigeria_content'].get('Progressive', 0),
                'Six selected creators; exploratory only.',
            ],
            [
                'Nigeria Audience',
                orientation_counts['nigeria_audience'].get('Regressive', 0),
                orientation_counts['nigeria_audience'].get('Progressive', 0),
                'Four selected creator-post files in current final workbook.',
            ],
        ],
    )
    doc.add_paragraph('Figure: Row balance by regressive and progressive creator groups')
    doc.add_picture(str(charts[1]), width=Inches(6.7))

    doc.add_heading('Country findings: Kenya (reported independently)', level=1)
    doc.add_paragraph(
        'Final workbook coverage: 394 content segments across 5 creators and 412 audience comments across 4 creator-post files.'
    )

    kenya_creator_rows = [[short_name_from_sheet(k), v] for k, v in kenya_content.items()]
    add_table(doc, ['Kenya content creator', 'Rows'], kenya_creator_rows)

    doc.add_paragraph('Exploratory LLM coding sample used for directional interpretation (Kenya content sample n = 241):')
    for topic, n in k_content_llm_stats['top_topics']:
        doc.add_paragraph(f'- {topic}: {n} mentions ({pct(n, k_content_llm_stats["n"])})')

    doc.add_paragraph('Exploratory audience sample (Kenya audience sample n = 200) shows the following stance distribution:')
    for k, v in k_audience_llm_stats['stance'].most_common():
        doc.add_paragraph(f'- {k}: {v} ({pct(v, k_audience_llm_stats["n"])})')

    kenya_figs = [
        BASE / 'Kenya/Audience Analysis Plots/07_funnel.png',
        BASE / 'Kenya/Audience Analysis Plots/09_orientation_compare.png',
    ]
    for p in kenya_figs:
        if p.exists():
            doc.add_paragraph(f'Figure: {p.name}')
            doc.add_picture(str(p), width=Inches(6.7))

    doc.add_heading('Country findings: Nigeria (reported independently)', level=1)
    doc.add_paragraph(
        'Final workbook coverage: 310 content segments across 6 creators and 417 audience comments across 4 creator-post files.'
    )

    nigeria_creator_rows = [[k, v] for k, v in nigeria_content.items()]
    add_table(doc, ['Nigeria content creator', 'Rows'], nigeria_creator_rows)

    doc.add_paragraph('Exploratory LLM coding sample used for directional interpretation (Nigeria content sample n = 260):')
    for topic, n in n_content_llm_stats['top_topics']:
        doc.add_paragraph(f'- {topic}: {n} mentions ({pct(n, n_content_llm_stats["n"])})')

    doc.add_paragraph('Exploratory audience sample (Nigeria audience sample n = 200) shows the following stance distribution:')
    for k, v in n_audience_llm_stats['stance'].most_common():
        doc.add_paragraph(f'- {k}: {v} ({pct(v, n_audience_llm_stats["n"])})')

    nigeria_figs = [
        BASE / 'Nigeria/Content Analysis/Exploratory/figures/content_theme_x_sentiment.png',
        BASE / 'Nigeria/Content Analysis/Exploratory/figures/content_theme_x_frame.png',
        BASE / 'Nigeria/Audience Analysis/Exploratory/figures/audience_theme_x_stance.png',
    ]
    for p in nigeria_figs:
        if p.exists():
            doc.add_paragraph(f'Figure: {p.name}')
            doc.add_picture(str(p), width=Inches(6.7))

    doc.add_heading('Engagement and reach layer', level=1)
    doc.add_paragraph(
        'Available metrics vary by platform and file. In current raw audience files, likes and replies are widely available, retweets are available primarily for X posts, and view/impression fields are largely missing. '
        'Accordingly, engagement analysis should use metrics wherever available and avoid forcing a single uniform metric across all sources.'
    )

    add_table(
        doc,
        ['Country', 'Raw audience files', 'Files with likes', 'Files with replies', 'Files with retweets', 'Files with views', 'Files with impressions'],
        [
            [
                'Kenya',
                kenya_raw_n,
                kenya_metric_presence.get('likes', 0),
                kenya_metric_presence.get('reply_count', 0),
                kenya_metric_presence.get('retweets', 0),
                kenya_metric_presence.get('views', 0),
                kenya_metric_presence.get('impressions', 0),
            ],
            [
                'Nigeria',
                nigeria_raw_n,
                nigeria_metric_presence.get('likes', 0),
                nigeria_metric_presence.get('reply_count', 0),
                nigeria_metric_presence.get('retweets', 0),
                nigeria_metric_presence.get('views', 0),
                nigeria_metric_presence.get('impressions', 0),
            ],
        ],
    )
    doc.add_paragraph('Figure: Engagement field availability in raw audience files')
    doc.add_picture(str(charts[2]), width=Inches(6.7))

    doc.add_paragraph('Reach context used for influencer selection (where metadata exists in current lists):')
    reach_rows = []
    for r in kenya_reach:
        if not r.get('found'):
            reach_rows.append(['Kenya', r['name'], 'Not found in current longlist metadata', '', '', '', ''])
        else:
            reach_rows.append([
                'Kenya',
                r['name'],
                r.get('orientation', ''),
                r.get('x', ''),
                r.get('instagram', ''),
                r.get('youtube', ''),
                r.get('reach_tier', ''),
            ])
    for r in nigeria_reach:
        if not r.get('found'):
            reach_rows.append(['Nigeria', r['name'], 'Not found in current longlist metadata', '', '', '', ''])
        else:
            reach_rows.append([
                'Nigeria',
                r['name'],
                r.get('orientation', ''),
                r.get('x', ''),
                r.get('instagram', ''),
                r.get('youtube', ''),
                r.get('reach_tier', ''),
            ])

    add_table(
        doc,
        ['Country', 'Influencer', 'Orientation in list', 'X reach', 'Instagram reach', 'YouTube reach', 'Reach tier'],
        reach_rows,
    )

    doc.add_heading('Limits and what this avoids claiming', level=1)
    doc.add_paragraph('- This revision avoids direct country ranking claims and uses country-by-country reporting.')
    doc.add_paragraph('- Percentages are tied to selected creators and sampled coding sets; they are not national prevalence estimates.')
    doc.add_paragraph('- View and impression metrics are incomplete for several platforms, so engagement modeling is limited to available fields.')

    doc.add_heading('Next focused-analysis steps', level=1)
    doc.add_paragraph('1. Build normalized engagement rates per post and per creator, using available likes/replies/retweets and denominator context where present.')
    doc.add_paragraph('2. Link engagement outcomes to coded narratives and framing variables within each country separately.')
    doc.add_paragraph('3. Report final focused results with clear denominator notes and platform-specific metric availability labels.')

    OUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOC)

    print(f'Wrote: {OUT_DOC}')
    print('Figure files:')
    for p in charts:
        print('-', p)


if __name__ == '__main__':
    main()
